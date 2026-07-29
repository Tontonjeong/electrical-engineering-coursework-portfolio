#!/usr/bin/env python3
"""Create a file-by-file source and publication audit for the coursework archive."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from PIL import Image
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from lxml import etree


PUBLIC_USED = "PUBLIC_USED"
PUBLIC_REFERENCE_ONLY = "PUBLIC_REFERENCE_ONLY"
DUPLICATE = "DUPLICATE"
PRIVATE_OR_CONFIDENTIAL = "PRIVATE_OR_CONFIDENTIAL"
THIRD_PARTY = "THIRD_PARTY"
CORRUPTED_OR_UNREADABLE = "CORRUPTED_OR_UNREADABLE"
INCOMPLETE_WORK = "INCOMPLETE_WORK"
NOT_RELEVANT = "NOT_RELEVANT_WITH_REASON"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def text_file(path: Path) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            return path.read_text(encoding=encoding), encoding
        except UnicodeDecodeError:
            continue
    raise UnicodeError("no supported text encoding")


def subject_and_relative(root: Path, path: Path) -> tuple[str, str]:
    parts = path.relative_to(root).parts
    if len(parts) >= 2:
        return parts[1], str(Path(*parts[2:]))
    return parts[0], ""


def is_generated(path_text: str, extension: str) -> bool:
    p = path_text.lower().replace("\\", "/")
    generated_segments = (
        "/.vs/",
        "/x64/debug/",
        ".cache/",
        "/xsim.dir/",
        "/copilotindices/",
        "/filecontentindex/",
        "/ipch/",
    )
    generated_extensions = {
        ".obj",
        ".pdb",
        ".ilk",
        ".idb",
        ".tlog",
        ".lastbuildstate",
        ".recipe",
        ".vsidx",
        ".vdb",
        ".rlx",
        ".reloc",
        ".rtti",
        ".svtype",
        ".type",
        ".xdbg",
        ".dbg",
        ".mem",
        ".pb",
        ".wdf",
        ".wpc",
        ".db",
        ".exe",
        ".c",
    }
    return any(token in p for token in generated_segments) or extension in generated_extensions


def canonical_score(path_text: str) -> tuple[int, int, int]:
    p = path_text.lower()
    return (
        1 if "__extracted" in p else 0,
        1 if p.endswith(".zip") else 0,
        len(path_text),
    )


def redact_public_path(value: str) -> str:
    """Remove direct personal identifiers while preserving file-level traceability."""
    student_id = "3221" + "1479"
    value = re.sub(rf"(?<!\d){student_id}(?!\d)", "[STUDENT-ID]", value)
    value = value.replace("류형록", "[AUTHOR]")
    return value


def classify(
    subject: str,
    relative: str,
    extension: str,
    duplicate: bool,
    extraction_error: bool,
) -> tuple[str, str]:
    p = relative.lower().replace("\\", "/")
    if "VRET" in subject.upper():
        return (
            PRIVATE_OR_CONFIDENTIAL,
            "Industry R&D material; withheld pending explicit disclosure review and redaction.",
        )
    if extraction_error:
        return CORRUPTED_OR_UNREADABLE, "Parser could not read the file structure."
    if duplicate:
        return DUPLICATE, "Exact SHA-256 duplicate of the canonical source record."
    if "ex12-3" in p or "시뮬하다가 중단" in p:
        return INCOMPLETE_WORK, "Source directory explicitly records an unfinished simulation."
    third_party_terms = (
        "cadence tutorial",
        "introduction to vhdl",
        "vhdl modeling",
        "4-bit full adder.pdf",
        "dec and mux",
        "finite state machinefsm",
        "universal shift register.pdf",
        "김동혁",
        "제11차 전력수급기본계획",
    )
    if any(term in p for term in third_party_terms):
        return THIRD_PARTY, "Lecture/tutorial/teammate/government reference; cite, do not republish."
    if is_generated(p, extension):
        return NOT_RELEVANT, "Generated IDE/simulator/compiler cache or binary build artifact."
    if extension == ".zip":
        return PUBLIC_REFERENCE_ONLY, "Archive provenance retained; publish curated extracted sources only."
    if extension in {".pwb", ".pwd", ".aux"}:
        return (
            PUBLIC_REFERENCE_ONLY,
            "Proprietary PowerWorld case retained for traceability; not directly web-publishable.",
        )
    if extension in {".xpr", ".wcfg", ".lpr", ".prj", ".tcl", ".bat", ".ini", ".log"}:
        return (
            PUBLIC_REFERENCE_ONLY,
            "Original project or simulator context retained as supporting execution evidence.",
        )
    if extension == ".png" and "사진/homework5/시뮬사진" in p:
        return (
            PUBLIC_REFERENCE_ONLY,
            "Large sequential Cadence screenshot set; select only decision-relevant frames.",
        )
    if extension in {
        ".pdf",
        ".docx",
        ".hwpx",
        ".xlsx",
        ".hwp",
        ".vhd",
        ".cpp",
        ".txt",
        ".png",
    }:
        return PUBLIC_USED, "Primary student-authored report, source, calculation, or visual evidence."
    return PUBLIC_REFERENCE_ONLY, "Retained for provenance; not selected for direct publication."


def dhash(path: Path) -> str:
    with Image.open(path) as image:
        gray = image.convert("L").resize((9, 8))
        pixels = list(gray.getdata())
    bits = []
    for row in range(8):
        offset = row * 9
        bits.extend(pixels[offset + col] > pixels[offset + col + 1] for col in range(8))
    value = sum((1 << index) for index, bit in enumerate(bits) if bit)
    return f"{value:016x}"


def inspect_content(path: Path, text_output: Path) -> dict[str, object]:
    suffix = path.suffix.lower()
    result: dict[str, object] = {
        "parser": "",
        "status": "metadata-only",
        "units": "",
        "unit_count": "",
        "text_chars": "",
        "formula_count": "",
        "table_count": "",
        "embedded_visual_count": "",
        "notes": "",
    }
    text = ""
    try:
        if suffix == ".pdf":
            reader = PdfReader(path)
            page_text = []
            image_count = 0
            for page in reader.pages:
                page_text.append(page.extract_text() or "")
                try:
                    image_count += len(page.images)
                except Exception:
                    pass
            text = "\n\n".join(page_text)
            result.update(
                parser="pypdf",
                status="parsed",
                units="pages",
                unit_count=len(reader.pages),
                text_chars=len(text),
                embedded_visual_count=image_count,
            )
        elif suffix == ".docx":
            document = Document(path)
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_text = []
            for table in document.tables:
                for row in table.rows:
                    table_text.append("\t".join(cell.text for cell in row.cells))
            text = "\n".join(paragraphs + table_text)
            with zipfile.ZipFile(path) as archive:
                media = [name for name in archive.namelist() if name.startswith("word/media/")]
            result.update(
                parser="python-docx",
                status="parsed",
                units="paragraphs",
                unit_count=len(document.paragraphs),
                text_chars=len(text),
                table_count=len(document.tables),
                embedded_visual_count=len(media),
            )
        elif suffix == ".hwpx":
            with zipfile.ZipFile(path) as archive:
                sections = sorted(
                    name
                    for name in archive.namelist()
                    if name.startswith("Contents/section") and name.endswith(".xml")
                )
                snippets = []
                for name in sections:
                    root = etree.fromstring(archive.read(name))
                    snippets.extend(root.xpath("//*[local-name()='t']/text()"))
                bins = [name for name in archive.namelist() if name.startswith("BinData/")]
            text = "\n".join(snippets)
            result.update(
                parser="HWPX XML",
                status="parsed",
                units="sections",
                unit_count=len(sections),
                text_chars=len(text),
                embedded_visual_count=len(bins),
            )
        elif suffix == ".xlsx":
            workbook = load_workbook(path, data_only=False, read_only=False)
            nonempty = 0
            formulas = 0
            charts = 0
            images = 0
            lines = []
            for sheet in workbook.worksheets:
                charts += len(sheet._charts)
                images += len(sheet._images)
                lines.append(f"## {sheet.title}")
                for row in sheet.iter_rows():
                    values = []
                    for cell in row:
                        if cell.value is not None:
                            nonempty += 1
                            if isinstance(cell.value, str) and cell.value.startswith("="):
                                formulas += 1
                            values.append(f"{cell.coordinate}={cell.value}")
                    if values:
                        lines.append("\t".join(values))
            text = "\n".join(lines)
            result.update(
                parser="openpyxl",
                status="parsed",
                units="worksheets",
                unit_count=len(workbook.sheetnames),
                text_chars=len(text),
                formula_count=formulas,
                table_count=nonempty,
                embedded_visual_count=charts + images,
                notes=f"nonempty_cells={nonempty}; charts={charts}; images={images}",
            )
        elif suffix in {".txt", ".vhd", ".cpp", ".log", ".tcl", ".bat", ".ini", ".aux"}:
            text, encoding = text_file(path)
            result.update(
                parser=f"text/{encoding}",
                status="parsed",
                units="lines",
                unit_count=text.count("\n") + 1,
                text_chars=len(text),
            )
        elif suffix == ".png":
            with Image.open(path) as image:
                result.update(
                    parser="Pillow",
                    status="parsed",
                    units="pixels",
                    unit_count=image.width * image.height,
                    embedded_visual_count=1,
                    notes=f"{image.width}x{image.height}; mode={image.mode}",
                )
        elif suffix == ".hwp":
            result.update(
                parser="none",
                status="unreadable",
                notes="Binary HWP requires a compatible HWP parser/application; withheld from claims.",
            )
        elif suffix in {".pwb", ".pwd"}:
            header = path.read_bytes()[:64]
            result.update(
                parser="binary signature",
                status="metadata-only",
                notes=f"header={header.hex()[:64]}",
            )
        else:
            result.update(parser="metadata", status="metadata-only")
    except Exception as exc:
        result.update(
            status="unreadable",
            notes=f"{type(exc).__name__}: {str(exc)[:240]}",
        )
    if text:
        out = text_output / f"{sha256(path)}.txt"
        if not out.exists():
            out.write_text(text, encoding="utf-8")
    return result


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fields is None:
        fields = list(rows[0]) if rows else []
    with path.open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-root", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--text-output", required=True)
    args = parser.parse_args()
    root = Path(args.source_root).resolve()
    output = Path(args.output).resolve()
    text_output = Path(args.text_output).resolve()
    output.mkdir(parents=True, exist_ok=True)
    text_output.mkdir(parents=True, exist_ok=True)

    files = sorted(path for path in root.rglob("*") if path.is_file())
    digests = {path: sha256(path) for path in files}
    groups: dict[str, list[Path]] = defaultdict(list)
    for path, digest in digests.items():
        groups[digest].append(path)
    canonicals = {
        digest: sorted(paths, key=lambda item: canonical_score(str(item.relative_to(root))))[0]
        for digest, paths in groups.items()
    }

    inventory: list[dict[str, object]] = []
    coverage: list[dict[str, object]] = []
    visuals: list[dict[str, object]] = []
    decisions: list[dict[str, object]] = []
    for index, path in enumerate(files, start=1):
        subject, relative = subject_and_relative(root, path)
        digest = digests[path]
        content = inspect_content(path, text_output)
        withheld = "VRET" in subject.upper()
        decision, reason = classify(
            subject,
            relative,
            path.suffix.lower(),
            path != canonicals[digest],
            content["status"] == "unreadable",
        )
        public_subject = "VRET (withheld)" if withheld else subject
        public_relative = (
            f"[WITHHELD]/VRET-{index:04d}{path.suffix.lower()}"
            if withheld
            else redact_public_path(relative)
        )
        public_digest = "[WITHHELD]" if withheld else digest
        public_canonical = (
            "[WITHHELD]"
            if withheld
            else redact_public_path(str(canonicals[digest].relative_to(root)))
        )
        public_content = (
            {
                "parser": "withheld",
                "status": "withheld",
                "units": "",
                "unit_count": "",
                "text_chars": "",
                "formula_count": "",
                "table_count": "",
                "embedded_visual_count": "",
                "notes": "Content metadata withheld pending disclosure review.",
            }
            if withheld
            else content
        )
        record = {
            "source_id": f"SRC-{index:04d}",
            "subject": public_subject,
            "relative_path": public_relative,
            "extension": path.suffix.lower() or "[none]",
            "bytes": path.stat().st_size,
            "sha256": public_digest,
            "exact_duplicate_count": len(groups[digest]),
            "canonical_relative_path": public_canonical,
            "classification": decision,
            "classification_reason": reason,
        }
        inventory.append(record)
        coverage.append(
            {
                "source_id": record["source_id"],
                "subject": public_subject,
                "relative_path": public_relative,
                **public_content,
            }
        )
        decisions.append(
            {
                "source_id": record["source_id"],
                "subject": public_subject,
                "relative_path": public_relative,
                "decision": decision,
                "public_action": {
                    PUBLIC_USED: "Curate into case study with source label/redaction.",
                    PUBLIC_REFERENCE_ONLY: "Retain in private audit; link only if useful and safe.",
                    DUPLICATE: "Do not republish duplicate bytes.",
                    PRIVATE_OR_CONFIDENTIAL: "Withhold; prepare redacted draft only after review.",
                    THIRD_PARTY: "Cite bibliographically; do not copy full asset.",
                    CORRUPTED_OR_UNREADABLE: "Do not claim content; document parser limitation.",
                    INCOMPLETE_WORK: "Show only as an explicitly incomplete engineering boundary.",
                    NOT_RELEVANT: "Exclude generated cache/build artifact from portfolio.",
                }[decision],
                "reason": reason,
            }
        )
        if path.suffix.lower() == ".png":
            with Image.open(path) as image:
                width, height = image.size
            role = "source screenshot"
            lower = relative.lower()
            if "파형" in lower or "wave" in lower:
                role = "waveform/result plot"
            elif "회로" in lower or "schematic" in lower or "design" in lower:
                role = "schematic/model"
            elif "응답" in lower or "그래" in lower or "sparameter" in lower:
                role = "quantitative result"
            visuals.append(
                {
                    "source_id": record["source_id"],
                    "subject": public_subject,
                    "relative_path": public_relative,
                    "width": width,
                    "height": height,
                    "aspect_ratio": round(width / height, 4) if height else "",
                    "sha256": public_digest,
                    "dhash": dhash(path),
                    "suggested_role": role,
                    "publication_decision": decision,
                    "notes": reason,
                }
            )

    duplicate_rows: list[dict[str, object]] = []
    group_index = 0
    for digest, paths in sorted(groups.items()):
        if len(paths) < 2:
            continue
        group_index += 1
        canonical = canonicals[digest]
        for path in sorted(paths):
            subject, relative = subject_and_relative(root, path)
            withheld = "VRET" in subject.upper()
            duplicate_rows.append(
                {
                    "duplicate_group": f"DUP-{group_index:04d}",
                    "sha256": "[WITHHELD]" if withheld else digest,
                    "bytes": path.stat().st_size,
                    "subject": "VRET (withheld)" if withheld else subject,
                    "relative_path": (
                        "[WITHHELD]" if withheld else redact_public_path(relative)
                    ),
                    "canonical": path == canonical,
                    "canonical_path": (
                        "[WITHHELD]"
                        if withheld
                        else redact_public_path(str(canonical.relative_to(root)))
                    ),
                }
            )

    write_csv(output / "source_inventory.csv", inventory)
    write_csv(output / "source_content_coverage.csv", coverage)
    write_csv(output / "source_visual_inventory.csv", visuals)
    write_csv(output / "duplicate_files.csv", duplicate_rows)
    write_csv(output / "publication_decisions.csv", decisions)

    class_counts = Counter(row["classification"] for row in inventory)
    subject_counts = Counter(row["subject"] for row in inventory)
    parsed = Counter(row["status"] for row in coverage)
    summary = {
        "files": len(files),
        "bytes": sum(path.stat().st_size for path in files),
        "subjects": dict(subject_counts),
        "classifications": dict(class_counts),
        "content_status": dict(parsed),
        "exact_duplicate_groups": sum(1 for paths in groups.values() if len(paths) > 1),
        "visual_files": len(visuals),
    }
    (output / "audit_summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (output / "final_coverage_report.md").write_text(
        "# Source Audit Coverage Report\n\n"
        f"- Total file records: **{summary['files']}**\n"
        f"- Expanded bytes: **{summary['bytes']:,}**\n"
        f"- Exact duplicate groups: **{summary['exact_duplicate_groups']}**\n"
        f"- Visual source files: **{summary['visual_files']}**\n"
        f"- Parsed content: **{parsed.get('parsed', 0)}**\n"
        f"- Metadata-only: **{parsed.get('metadata-only', 0)}**\n"
        f"- Unreadable: **{parsed.get('unreadable', 0)}**\n\n"
        "## Publication classification\n\n"
        + "\n".join(f"- {key}: **{value}**" for key, value in sorted(class_counts.items()))
        + "\n\nEvery extracted file has a source ID, SHA-256, content-inspection status, "
        "publication decision, and reason. Exact duplicates retain one canonical record. "
        "VRET material remains withheld pending a separate disclosure review.\n",
        encoding="utf-8",
    )
    print(json.dumps(summary, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
